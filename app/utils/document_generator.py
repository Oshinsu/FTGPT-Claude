from typing import Dict, Any
from datetime import datetime
from pathlib import Path
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
import pdfkit


class DocumentGenerator:
    """Générateur de documents (CV, lettres de motivation)."""
    
    def __init__(self):
        self.output_dir = Path("generated_documents")
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_cv(self, data: Dict[str, Any]) -> Path:
        """Génère un CV au format DOCX."""
        doc = Document()
        
        # En-tête avec informations personnelles
        header = doc.add_heading(data.get("name", "Nom Prénom"), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations de contact
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.add_run(f"{data.get('email', '')} | {data.get('phone', '')}")
        if data.get("address"):
            contact_para.add_run(f"\n{data['address']}")
        if data.get("linkedin"):
            contact_para.add_run(f"\n{data['linkedin']}")
        
        # Objectif professionnel
        if data.get("target_job"):
            doc.add_heading("Objectif Professionnel", 1)
            doc.add_paragraph(f"Recherche un poste de {data['target_job']}")
        
        # Expériences professionnelles
        if data.get("experiences"):
            doc.add_heading("Expériences Professionnelles", 1)
            for exp in data["experiences"]:
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}").bold = True
                p.add_run(f"\n{exp.get('period', '')}")
                
                # Missions
                if exp.get("missions"):
                    for mission in exp["missions"]:
                        doc.add_paragraph(f"• {mission}", style='List Bullet')
        
        # Compétences
        if data.get("skills"):
            doc.add_heading("Compétences", 1)
            skills_para = doc.add_paragraph()
            
            if isinstance(data["skills"], list):
                skills_text = " • ".join(data["skills"])
            else:
                skills_text = data["skills"]
            
            skills_para.add_run(skills_text)
        
        # Formation
        if data.get("education"):
            doc.add_heading("Formation", 1)
            for edu in data["education"]:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')} - {edu.get('school', '')}").bold = True
                p.add_run(f"\n{edu.get('year', '')}")
        
        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"CV_{data.get('name', 'candidat').replace(' ', '_')}_{timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        
        return filepath
    
    def generate_cover_letter(self, data: Dict[str, Any]) -> Path:
        """Génère une lettre de motivation."""
        doc = Document()
        
        # En-tête expéditeur
        doc.add_paragraph(data.get("name", "Nom Prénom"))
        doc.add_paragraph(data.get("address", "Adresse"))
        doc.add_paragraph(data.get("email", "email@example.com"))
        doc.add_paragraph(data.get("phone", "Téléphone"))
        
        doc.add_paragraph()
        
        # Destinataire
        doc.add_paragraph(data.get("company_name", "Nom de l'entreprise"))
        doc.add_paragraph(data.get("company_address", "Adresse de l'entreprise"))
        
        doc.add_paragraph()
        
        # Lieu et date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"Paris, le {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        
        # Objet
        doc.add_paragraph(f"Objet : {data.get('object', 'Candidature')}")
        
        doc.add_paragraph()
        
        # Corps de la lettre
        doc.add_paragraph("Madame, Monsieur,")
        
        doc.add_paragraph()
        
        # Paragraphes de contenu
        if data.get("content"):
            for paragraph in data["content"]:
                doc.add_paragraph(paragraph)
                doc.add_paragraph()
        
        # Formule de politesse
        doc.add_paragraph(
            "Dans l'attente de votre réponse, je vous prie d'agréer, "
            "Madame, Monsieur, l'expression de mes salutations distinguées."
        )
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature.add_run(data.get("name", "Nom Prénom"))
        
        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"LM_{data.get('company_name', 'entreprise').replace(' ', '_')}_{timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)
        
        return filepath
    
    def save_document(self, content: str, doc_type: str) -> Path:
        """Sauvegarde un document généré."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{doc_type}_{timestamp}.md"
        filepath = self.output_dir / filename
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        
        return filepath
